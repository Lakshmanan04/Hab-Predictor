# app.py
import os
import io
import time
from typing import Tuple, List

import numpy as np
import pandas as pd
from PIL import Image
import matplotlib.pyplot as plt

import streamlit as st

import torch
import torch.nn as nn
import torch.nn.functional as F
from torchvision import transforms

# Interpretability libs
from pytorch_grad_cam import GradCAM
from pytorch_grad_cam.utils.model_targets import ClassifierOutputTarget

# Load scaler alongside model
import joblib
scaler_features = joblib.load('scaler_features.save')
scaler_cellcount = joblib.load('scaler_cellcount.save')

import requests
import base64
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


# ---------------------------
# Model definitions
# ---------------------------

class PatchEmbed(nn.Module):
    def __init__(self, in_ch=3, patch_size=16, emb_dim=256, img_size=128):
        super().__init__()
        self.patch_size = patch_size
        self.n_patches = (img_size // patch_size) ** 2
        self.proj = nn.Conv2d(in_ch, emb_dim, kernel_size=patch_size, stride=patch_size)
        self.cls_token = nn.Parameter(torch.zeros(1, 1, emb_dim))
        self.pos_emb = nn.Parameter(torch.randn(1, self.n_patches + 1, emb_dim))

    def forward(self, x):
        x = self.proj(x)
        x = x.flatten(2).transpose(1, 2)
        b = x.size(0)
        cls = self.cls_token.expand(b, -1, -1)
        x = torch.cat([cls, x], dim=1) + self.pos_emb
        return x

class ViTEncoder(nn.Module):
    def __init__(self, emb_dim=256, n_layers=6, n_heads=8, mlp_dim=512, dropout=0.1):
        super().__init__()
        layer = nn.TransformerEncoderLayer(
            d_model=emb_dim, nhead=n_heads,
            dim_feedforward=mlp_dim,
            dropout=dropout, activation='gelu',
            batch_first=True)
        self.encoder = nn.TransformerEncoder(layer, num_layers=n_layers)

    def forward(self, x):
        return self.encoder(x)

class TemporalTransformer(nn.Module):
    def __init__(self, input_dim=3, emb_dim=128, n_layers=3, n_heads=4, mlp_dim=256, max_len=1):
        super().__init__()
        self.input_proj = nn.Linear(input_dim, emb_dim)
        self.pos_emb = nn.Parameter(torch.randn(1, max_len, emb_dim))
        layer = nn.TransformerEncoderLayer(
            d_model=emb_dim, nhead=n_heads,
            dim_feedforward=mlp_dim,
            activation='relu', batch_first=True)
        self.encoder = nn.TransformerEncoder(layer, num_layers=n_layers)

    def forward(self, seq):
        if seq.dim() == 2:
            seq = seq.unsqueeze(1)
        b, t, _ = seq.shape
        x = self.input_proj(seq) + self.pos_emb[:, :t, :]
        x = self.encoder(x)
        pooled = x.mean(dim=1)
        return x, pooled

class DualTransformerFusion(nn.Module):
    def __init__(self, img_emb=256, patch_size=16, img_size=128, seq_in=3, seq_emb=128, num_classes=4):
        super().__init__()
        self.patch = PatchEmbed(in_ch=3, patch_size=patch_size, emb_dim=img_emb, img_size=img_size)
        self.vit = ViTEncoder(emb_dim=img_emb, n_layers=6, n_heads=8, mlp_dim=img_emb * 2)
        self.temporal = TemporalTransformer(input_dim=seq_in, emb_dim=seq_emb, n_layers=4, n_heads=4, mlp_dim=seq_emb * 2)
        self.temporal_proj = nn.Linear(seq_emb, img_emb)
        self.cross_attn = nn.MultiheadAttention(embed_dim=img_emb, num_heads=8, batch_first=True)
        self.classifier = nn.Sequential(
            nn.Linear(img_emb * 2, 256),
            nn.ReLU(),
            nn.Dropout(0.3),
            nn.Linear(256, num_classes)
        )

    def forward(self, image, seq):
        img_tokens = self.patch(image)
        img_tokens = self.vit(img_tokens)
        seq_out, seq_pooled = self.temporal(seq)
        q = self.temporal_proj(seq_pooled).unsqueeze(1)
        attn_out, _ = self.cross_attn(q, img_tokens, img_tokens)
        img_cls = img_tokens[:, 0, :]
        fused = torch.cat([img_cls, attn_out.squeeze(1)], dim=1)
        return self.classifier(fused)

class ImageOnlyModel(nn.Module):
    def __init__(self, full_model, n_classes=4):
        super().__init__()
        self.patch = full_model.patch
        self.vit = full_model.vit
        self.head = nn.Linear(self.patch.proj.out_channels, n_classes)

    def forward(self, image):
        img_tokens = self.patch(image)
        img_tokens = self.vit(img_tokens)
        img_cls = img_tokens[:, 0, :]
        return self.head(img_cls)

# ---------------------------
# Preprocessing & utilities
# ---------------------------
IMG_SIZE = 128
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

transform = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.ToTensor(),
])

CLASS_NAMES = ["No Bloom", "Low", "Moderate", "High"]

SEVERITY_COLORS = {
    "No Bloom": "#22c55e",      # Green for no bloom
    "Low": "#10b981",
    "Moderate": "#f59e0b",
    "High": "#ef4444",
}

SEVERITY_EMOJIS = {
    "No Bloom": "✅",
    "Low": "🟢",
    "Moderate": "🟠",
    "High": "🔴",
}

@st.cache_resource(show_spinner=False)
def load_trained_model(checkpoint_path: str = "checkpoints/final_model.pth"):
    model = DualTransformerFusion().to(device)
    if not os.path.exists(checkpoint_path):
        st.warning(f"⚠️ Checkpoint not found: {checkpoint_path}")
        model.eval()
        image_model = ImageOnlyModel(model).to(device)
        image_model.eval()
        return model, image_model, None

    ckpt = torch.load(checkpoint_path, map_location=device)
    model.load_state_dict(ckpt['model_state_dict'])
    model.eval()
    image_model = ImageOnlyModel(model).to(device)
    image_model.eval()
    target_layers = [image_model.patch.proj]
    return model, image_model, target_layers

def preprocess_image_pil(img_pil: Image.Image) -> torch.Tensor:
    img = img_pil.convert("RGB")
    return transform(img).unsqueeze(0)

def preprocess_sensors(salinity: float, water_temp: float, cellcount: float) -> torch.Tensor:
    # Scale salinity and water_temp
    features_scaled = scaler_features.transform(np.array([[salinity, water_temp]], dtype=np.float32))
    
    # Log-transform and scale cellcount separately
    cellcount_log = np.log1p(cellcount)
    cellcount_scaled = scaler_cellcount.transform(np.array([[cellcount_log]], dtype=np.float32))
    
    # Combine features back into one array
    combined = np.hstack([features_scaled, cellcount_scaled])
    
    # Convert to tensor and move to device
    return torch.tensor(combined, dtype=torch.float32).to(device)


def predict(model: DualTransformerFusion, image_tensor: torch.Tensor, sensor_tensor: torch.Tensor) -> Tuple[np.ndarray, np.ndarray]:
    with torch.no_grad():
        image_tensor = image_tensor.to(device)
        sensor_tensor = sensor_tensor.to(device)
        logits = model(image_tensor, sensor_tensor)
        probs = F.softmax(logits, dim=1).cpu().numpy()[0]
        pred_idx = int(np.argmax(probs))
        return pred_idx, probs

def compute_gradcam(image_model: ImageOnlyModel, target_layers: List[nn.Module], image_tensor: torch.Tensor, class_idx: int):
    image_model.eval()
    # Fix: remove use_cuda parameter
    cam = GradCAM(model=image_model, target_layers=target_layers)
    targets = [ClassifierOutputTarget(class_idx)]
    grayscale_cam = cam(input_tensor=image_tensor.to(device), targets=targets, eigen_smooth=True, aug_smooth=True)
    cam_map = grayscale_cam[0]
    cam_map = np.nan_to_num(cam_map)
    cam_map = np.clip(cam_map, 0.0, 1.0)
    return cam_map

def overlay_cam_on_image(img_pil: Image.Image, cam_map: np.ndarray, alpha=0.5) -> Image.Image:
    img = img_pil.convert("RGB").resize((cam_map.shape[1], cam_map.shape[0]))
    img_np = np.array(img)
    img_norm = (img_np - img_np.min()) / (img_np.max() - img_np.min() + 1e-8)
    img_uint8 = (img_norm * 255).astype(np.uint8)
    heatmap = (cam_map * 255).astype(np.uint8)
    heatmap_bgr = plt.cm.jet(heatmap)[:, :, :3]
    heatmap_rgb = (heatmap_bgr * 255).astype(np.uint8)
    overlay = (alpha * heatmap_rgb + (1 - alpha) * img_uint8).astype(np.uint8)
    return Image.fromarray(overlay)

# Initialize Gemini client (ensure GOOGLE_API_KEY is set in env variables)
GEMINI_API_KEY = st.secrets["GEMINIAPIKEY"]
GENAI_MODEL = "gemini-flash-latest"  # or set via env

def generate_report(prompt_text, max_output_tokens=5000, temperature=0.7):
    """
    Calls Gemini Generative Language API synchronously,
    returns generated report text or error string.
    """

    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not set in environment. Put it in .env")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GENAI_MODEL}:generateContent?key={GEMINI_API_KEY}"

    request_json = {
        "contents": [
            {
                "parts": [
                    {"text": prompt_text}
                ]
            }
        ],
        "generationConfig": {
            "temperature": temperature,
            "maxOutputTokens": max_output_tokens
        }
    }

    headers = {"Content-Type": "application/json; charset=utf-8"}

    try:
        resp = requests.post(url, json=request_json, headers=headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        return f"[GenAI error or network issue: {e}]"

def encode_image_base64(img: Image.Image) -> str:
    buff = io.BytesIO()
    img.save(buff, format='PNG')
    encoded = base64.b64encode(buff.getvalue()).decode('utf-8')
    return encoded

def create_word_report(report_text, pred_name, pred_prob, salinity, water_temp, cellcount):
    """
    Creates a Word document from the report text with proper formatting.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('HAB Severity Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'Generated: {time.strftime("%B %d, %Y %H:%M:%S")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Summary Section
    doc.add_heading('Executive Summary', 1)
    summary = doc.add_paragraph()
    summary.add_run(f'Predicted Severity: ').bold = True
    summary.add_run(f'{pred_name}\n')
    summary.add_run(f'Confidence Level: ').bold = True
    summary.add_run(f'{pred_prob*100:.1f}%\n')
    
    # Sensor Data Section
    doc.add_heading('Sensor Data', 1)
    sensor_table = doc.add_table(rows=4, cols=2)
    sensor_table.style = 'Light Grid Accent 1'
    
    sensor_data = [
        ('Parameter', 'Value'),
        ('Salinity', f'{salinity} PSU'),
        ('Water Temperature', f'{water_temp}°C'),
        ('Cell Count', f'{cellcount}')
    ]
    
    for i, (param, val) in enumerate(sensor_data):
        row = sensor_table.rows[i]
        row.cells[0].text = param
        row.cells[1].text = str(val)
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Detailed Analysis Section
    doc.add_heading('Detailed Analysis', 1)
    
    # Split report text into paragraphs and add them
    paragraphs = report_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Check if it's a heading (starts with #)
            if para.strip().startswith('#'):
                heading_text = para.strip().lstrip('#').strip()
                doc.add_heading(heading_text, 2)
            else:
                doc.add_paragraph(para.strip())
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Generated by HAB Severity Classifier | AI-Powered Environmental Monitoring')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# ---------------------------
# Custom CSS
# ---------------------------
def load_custom_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    * {
        font-family: 'Inter', sans-serif;
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    .main {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        padding: 0 !important;
    }
    
    .block-container {
        padding: 2rem 3rem !important;
        max-width: 1400px !important;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background: #1e293b;
        border-right: 1px solid #334155;
    }
    
    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] {
        color: #e2e8f0;
    }
    
    /* Navigation tile styling */
    .nav-tile {
        background: #334155;
        border: 2px solid #475569;
        border-radius: 12px;
        padding: 1.2rem;
        margin: 0.8rem 0;
        cursor: pointer;
        transition: all 0.3s ease;
        text-align: center;
        color: #cbd5e1;
        font-weight: 500;
        font-size: 1.05rem;
    }
    
    .nav-tile:hover {
        background: #475569;
        border-color: #667eea;
        transform: translateX(4px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.3);
    }
    
    .nav-tile.active {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-color: #667eea;
        color: white;
        font-weight: 600;
    }
    
    /* Header card */
    .main-header {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        border: 1px solid #475569;
        padding: 2rem;
        border-radius: 20px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
        margin-bottom: 2rem;
    }
    
    .main-title {
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 0;
        padding: 0;
    }
    
    .subtitle {
        color: #94a3b8;
        font-size: 1.1rem;
        margin-top: 0.5rem;
    }
    
    /* Card styling */
    .custom-card {
        background: #1e293b;
        border: 1px solid #334155;
        border-radius: 16px;
        padding: 1.5rem;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3);
        margin-bottom: 1.5rem;
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }
    
    .custom-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 15px 40px rgba(102, 126, 234, 0.2);
        border-color: #475569;
    }
    
    .custom-card h3, .custom-card h4, .custom-card h5 {
        color: #f1f5f9 !important;
        margin-top: 0;
    }
    
    .custom-card p, .custom-card li {
        color: #cbd5e1;
    }
    
    /* Result card */
    .result-card {
        background: #1e293b;
        border: 1px solid #334155;
        border-radius: 20px;
        padding: 2rem;
        position: relative;
        overflow: hidden;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
        margin: 2rem 0;
    }
    
    .result-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 6px;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
    }
    
    .result-card h2 {
        color: #f1f5f9;
        margin-top: 0;
    }
    
    /* Severity badge */
    .severity-badge {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.75rem 1.5rem;
        border-radius: 50px;
        font-weight: 600;
        font-size: 1.25rem;
        margin: 1rem 0;
        animation: fadeInUp 0.5s ease;
    }
    
    @keyframes fadeInUp {
        from {
            opacity: 0;
            transform: translateY(20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    /* Confidence meter */
    .confidence-bar {
        height: 12px;
        background: #334155;
        border-radius: 10px;
        overflow: hidden;
        position: relative;
    }
    
    .confidence-fill {
        height: 100%;
        border-radius: 10px;
        transition: width 1s ease;
        position: relative;
        overflow: hidden;
    }
    
    .confidence-fill::after {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        bottom: 0;
        right: 0;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
        animation: shimmer 2s infinite;
    }
    
    @keyframes shimmer {
        0% { transform: translateX(-100%); }
        100% { transform: translateX(100%); }
    }
    
    /* Button styling */
    .stButton > button {
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 1rem 2rem;
        font-size: 1.1rem;
        font-weight: 600;
        border-radius: 12px;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
    }
    
    /* Input fields */
    .stNumberInput > div > div > input {
        background: #334155 !important;
        border: 1px solid #475569 !important;
        color: #f1f5f9 !important;
        border-radius: 8px !important;
    }
    
    .stNumberInput label {
        color: #cbd5e1 !important;
        font-weight: 500 !important;
    }
    
    /* File uploader */
    [data-testid="stFileUploader"] {
        background: #334155;
        border: 2px dashed #667eea;
        border-radius: 12px;
        padding: 2rem;
    }
    
    [data-testid="stFileUploader"] label {
        color: #cbd5e1 !important;
    }
    
    /* Metric cards */
    .metric-card {
        background: #334155;
        border: 1px solid #475569;
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: 700;
        color: #f1f5f9;
        margin: 0.5rem 0;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: #94a3b8;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    
    /* Info boxes */
    .info-box {
        background: rgba(102, 126, 234, 0.15);
        border-left: 4px solid #667eea;
        padding: 1.5rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #cbd5e1;
    }
    
    /* Text colors */
    .stMarkdown {
        color: #e2e8f0;
    }
    
    h1, h2, h3, h4, h5, h6 {
        color: #f1f5f9 !important;
    }
    
    p {
        color: #cbd5e1;
    }
    
    /* Success, warning, error, info */
    .stSuccess, .stWarning, .stError, .stInfo {
        background: #1e293b !important;
        border: 1px solid #334155 !important;
        color: #cbd5e1 !important;
    }
    
    /* Section divider */
    .section-divider {
        height: 1px;
        background: linear-gradient(90deg, transparent, #475569, transparent);
        margin: 3rem 0;
    }
    </style>
    """, unsafe_allow_html=True)

# ---------------------------
# Streamlit UI
# ---------------------------
st.set_page_config(
    page_title="HAB Severity Classifier",
    page_icon="🌊",
    layout="wide",
    initial_sidebar_state="expanded"
)

load_custom_css()

# Initialize session state for navigation
if 'page' not in st.session_state:
    st.session_state.page = 'home'

# Minimal Sidebar with tile navigation
with st.sidebar:
    st.markdown("### 🌊 HAB Inspector")
    st.markdown("---")
    
    # Custom tile navigation
    if st.button("🏠 Home & Analysis", key="nav_home", use_container_width=True):
        st.session_state.page = 'home'
    
    if st.button("🔬 Model Architecture", key="nav_arch", use_container_width=True):
        st.session_state.page = 'arch'

# Load model
MODEL_PATH = os.environ.get("HAB_CHECKPOINT", "checkpoints/final_model.pth")
full_model, image_model, target_layers = load_trained_model(MODEL_PATH)

# HOME PAGE (Combined with About HABs)
if st.session_state.page == 'home':
    st.markdown("""
    <div class="main-header">
        <h1 class="main-title">🌊 HAB Severity Classifier</h1>
        <p class="subtitle">AI-Powered Harmful Algal Bloom Detection & Analysis</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Analysis Section
    col1, col2 = st.columns([1, 1.5], gap="large")
    
    with col1:
        st.markdown("### 📤 Upload Water Imagery")
        uploaded_file = st.file_uploader(
            "Drag and drop or browse files",
            type=["png", "jpg", "jpeg", "tif"],
            help="Upload Sentinel-2 imagery or water surface photos",
            label_visibility="collapsed"
        )
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("### 🎛️ Sensor Readings")
        
        col_a, col_b = st.columns(2)
        with col_a:
            salinity = st.number_input(
                "💧 Salinity (PSU)",
                value=35.0,
                format="%.2f"
            )
            cellcount = st.number_input(
                "🔬 Cell Count",
                value=1000.0,
                format="%.0f"
            )
        
        with col_b:
            water_temp = st.number_input(
                "🌡️ Temperature (°C)",
                value=25.0,
                format="%.1f"
            )
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        run_btn = st.button("🔍 Analyze HAB Severity", use_container_width=True)
    
    with col2:
        if run_btn:
            if uploaded_file is None:
                st.warning("⚠️ Please upload an image before analyzing.")
            else:
                with st.spinner("🔄 Analyzing water sample..."):
                    time.sleep(0.5)
                    try:
                        img_pil = Image.open(uploaded_file).convert("RGB")
                        image_tensor = preprocess_image_pil(img_pil)
                        sensor_tensor = preprocess_sensors(salinity, water_temp, cellcount)

                        pred_idx, probs = predict(full_model, image_tensor, sensor_tensor)
                        pred_name = CLASS_NAMES[pred_idx]
                        pred_prob = float(probs[pred_idx])
                        
                        # Store in session state for report
                        st.session_state['pred_name'] = pred_name
                        st.session_state['pred_prob'] = pred_prob
                        st.session_state['salinity'] = salinity
                        st.session_state['water_temp'] = water_temp
                        st.session_state['cellcount'] = cellcount
                        color = SEVERITY_COLORS[pred_name]
                        emoji = SEVERITY_EMOJIS[pred_name]

                        st.markdown(f"""
                        <div class="result-card">
                            <h2>Analysis Results</h2>
                            <div class="severity-badge" style="background: {color}30; color: {color}; border: 2px solid {color};">
                                <span style="font-size: 1.5rem;">{emoji}</span>
                                <span>HAB Severity: <strong>{pred_name}</strong></span>
                            </div>
                            <div style="margin: 1.5rem 0;">
                                <div style="display: flex; justify-content: space-between; margin-bottom: 0.5rem;">
                                    <span style="font-weight: 600; color: #94a3b8;">Confidence Level</span>
                                    <span style="font-weight: 700; color: {color};">{pred_prob*100:.1f}%</span>
                                </div>
                                <div class="confidence-bar">
                                    <div class="confidence-fill" style="width: {pred_prob*100}%; background: {color};"></div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        # GradCAM Visualization
                        st.markdown("### 🔍 Visual Analysis")

                        if target_layers is not None:
                            try:
                                cam_map = compute_gradcam(image_model, target_layers, image_tensor, pred_idx)
                                overlay = overlay_cam_on_image(img_pil, cam_map)

                                img_col1, img_col2 = st.columns(2)
                                with img_col1:
                                    st.image(img_pil, caption="Original Image", use_container_width=True)
                                with img_col2:
                                    st.image(overlay, caption="Attention Heatmap", use_container_width=True)

                                st.info("🔥 Heatmap highlights image regions that influenced the prediction")
                            except Exception as e:
                                st.error(f"Visualization failed: {str(e)}")
                        else:
                            st.warning("⚠️ GradCAM unavailable - checkpoint not loaded")

                        # Gemini report generation prompt
                        # Before calling generate_report_async(prompt), encode images:
                        heatmap_b64 = encode_image_base64(overlay)
                        original_b64 = encode_image_base64(img_pil)  # the original input image

                        # Modify your prompt to include images as base64 data URLs and instructions
                        prompt = f"""
                        You are an environmental expert reviewing water sample analysis with visual heatmaps.

                        Original water image:
                        ![original](data:image/png;base64,{original_b64})

                        Heatmap overlay on the water image:
                        ![heatmap](data:image/png;base64,{heatmap_b64})

                        Based on the above images, prediction: {pred_name} ({pred_prob*100:.1f}% confidence)
                        Sensor data: Salinity={salinity} PSU, Temp={water_temp}°C, Cell Count={cellcount}

                        Please provide a detailed report explaining the prediction, highlighting the regions of interest, environmental implications, and recommended actions.
                        """


                        # Synchronous Gemini report generation
                        report_text = generate_report(prompt)

                        st.session_state['generated_report'] = report_text
                        st.session_state['report_ready'] = True

                    except Exception as e:
                        st.error(f"❌ Analysis failed: {str(e)}")
        
        else:
            st.markdown("""
            <div style="text-align: center; padding: 4rem 2rem; background: #1e293b; border: 1px solid #334155; border-radius: 20px; box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3);">
                <h2 style="color: #94a3b8; margin-bottom: 1rem;">🌊 Ready to Analyze</h2>
                <p style="color: #64748b; font-size: 1.1rem;">Upload an image and enter sensor data to begin</p>
                <div style="margin-top: 2rem; display: flex; justify-content: center; gap: 2rem;">
                    <div style="text-align: center;">
                        <div style="font-size: 2rem; margin-bottom: 0.5rem;">📸</div>
                        <div style="color: #94a3b8; font-weight: 500;">Upload Image</div>
                    </div>
                    <div style="text-align: center;">
                        <div style="font-size: 2rem; margin-bottom: 0.5rem;">🎛️</div>
                        <div style="color: #94a3b8; font-weight: 500;">Enter Sensors</div>
                    </div>
                    <div style="text-align: center;">
                        <div style="font-size: 2rem; margin-bottom: 0.5rem;">🔍</div>
                        <div style="color: #94a3b8; font-weight: 500;">Analyze</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        if 'report_ready' in st.session_state and st.session_state['report_ready']:
            if st.button("📄 View Detailed Report", key="view_report_btn"):
                st.session_state.page = 'report'
                st.rerun()
        elif 'report_ready' in st.session_state and not st.session_state['report_ready']:
            st.info("⏳ Generating detailed report, please wait...")
    
    # Section divider
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
    
    # About HABs Section
    st.markdown("""
    <div class="main-header">
        <h2 class="main-title" style="font-size: 2rem;">📊 Understanding Harmful Algal Blooms</h2>
        <p class="subtitle">Environmental Impact & Detection Importance</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 🌊 What are HABs?")
    st.markdown("""
    **Harmful Algal Blooms** are rapid Algal population increases that produce toxins 
    or cause environmental damage. Climate change and pollution are increasing their frequency.
    """)
    st.markdown('</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⚠️ Causes")
        st.markdown("""
        - **Nutrient Pollution**: Nitrogen & phosphorus runoff
        - **Warm Temperatures**: Climate change effects
        - **Stagnant Water**: Low flow conditions
        - **Light Availability**: Shallow, clear waters
        - **Water Stability**: Reduced mixing
        """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("#### 💀 Impacts")
        st.markdown("""
        - **Toxins**: Harmful to humans and wildlife
        - **Oxygen Depletion**: Creates hypoxic zones
        - **Light Blockage**: Prevents vegetation growth
        - **Food Web**: Disrupts ecosystem balance
        - **Economics**: Damages tourism and fishing
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown("#### 🎯 Detection Benefits")
        st.markdown("""
        - **Public Health**: Prevent toxic exposure
        - **Water Treatment**: Adjust filtration processes
        - **Fisheries**: Protect aquaculture operations
        - **Ecosystem**: Monitor environmental health
        - **Policy**: Guide nutrient management
        """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("#### 🔬 Our Solution")
        st.markdown("""
        - **Satellite Imagery**: Remote sensing capability
        - **Sensor Networks**: Real-time water quality data
        - **AI Models**: Automated severity classification
        - **Interpretability**: Explainable AI insights
        - **Rapid Response**: Quick assessment for action
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("### 📈 Severity Classifications")
    
    severity_cols = st.columns(4)
    
    severity_descriptions = {
        "No Bloom": "No bloom detected",
        "Low": "Minimal concern",
        "Moderate": "Watch status",
        "High": "Action needed"
    }

    for i, (cls_name, color, emoji) in enumerate(zip(CLASS_NAMES, 
                                                      [SEVERITY_COLORS[c] for c in CLASS_NAMES],
                                                      [SEVERITY_EMOJIS[c] for c in CLASS_NAMES])):
        with severity_cols[i]:
            st.markdown(f"""
            <div style="background: {color}20; border: 2px solid {color}; border-radius: 12px; padding: 1.5rem; text-align: center;">
                <div style="font-size: 3rem; margin-bottom: 0.5rem;">{emoji}</div>
                <div style="font-weight: 700; font-size: 1.2rem; color: {color}; margin-bottom: 0.5rem;">{cls_name}</div>
                <div style="color: #94a3b8; font-size: 0.85rem;">
                    {severity_descriptions[cls_name]}
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="info-box">', unsafe_allow_html=True)
    st.markdown("""
    💡 **Learn More**: For additional resources, visit [NOAA HAB Portal](https://oceanservice.noaa.gov/hazards/hab/) 
    or [EPA Water Quality Portal](https://www.epa.gov/waterdata)
    """)
    st.markdown('</div>', unsafe_allow_html=True)

# ARCHITECTURE PAGE
elif st.session_state.page == 'arch':
    st.markdown("""
    <div class="main-header">
        <h1 class="main-title">🔬 Model Architecture</h1>
        <p class="subtitle">Dual-Transformer Fusion Network for Multimodal HAB Detection</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 🏗️ Architecture Overview")
    
    arch_img_path = "architecture.jpg"
    if os.path.exists(arch_img_path):
        st.image(arch_img_path, use_container_width=True)
    else:
        st.info("📁 Architecture diagram not found at `architecture.jpg`")
    
    st.markdown("""
    The model combines two transformer-based encoders to process multimodal inputs for accurate HAB severity prediction.
    """)
    st.markdown('</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🖼️ Vision Transformer (ViT)")
        st.markdown("""
        - **Patch Embedding**: Splits images into 16×16 patches
        - **Transformer Encoder**: 6 layers with 8 attention heads
        - **Output**: Rich visual feature representations
        - **Purpose**: Captures spatial patterns in water imagery
        """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("#### 🎯 Cross-Attention Fusion")
        st.markdown("""
        - Temporal features query visual tokens
        - Multi-head attention mechanism (8 heads)
        - Learns which image regions matter most
        - Produces context-aware visual features
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown("#### ⏱️ Temporal Transformer")
        st.markdown("""
        - **Input**: Sensor readings (salinity, temperature, cell count)
        - **Encoder**: 4 layers with 4 attention heads
        - **Output**: Contextual sensor embeddings
        - **Purpose**: Models environmental factor relationships
        """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("#### 🎲 Classification Head")
        st.markdown("""
        - Concatenates visual CLS token + attended features
        - 2-layer MLP with dropout regularization (0.3)
        - Outputs 4-class severity predictions
        - Uses softmax for probability distribution
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("### 🔍 Interpretability Method")
    
    st.markdown("""
    **GradCAM (Gradient-weighted Class Activation Mapping)**
    
    GradCAM provides visual explanations by highlighting which regions of the input image 
    were most important for the model's prediction. It works by:
    
    - Computing gradients of the predicted class with respect to feature maps
    - Generating a coarse localization map highlighting important regions
    - Overlaying this heatmap on the original image
    - Showing where the model "looks" when making decisions
    
    This helps users understand and trust the model's predictions by providing transparent, 
    interpretable visualizations of the decision-making process.
    """)
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("### 📊 Model Specifications")
    
    spec_col1, spec_col2, spec_col3 = st.columns(3)
    
    with spec_col1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Image Size</div>
            <div class="metric-value">{IMG_SIZE}×{IMG_SIZE}</div>
            <div style="color: #94a3b8; font-size: 0.85rem;">pixels</div>
        </div>
        """, unsafe_allow_html=True)
    
    with spec_col2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Classes</div>
            <div class="metric-value">{len(CLASS_NAMES)}</div>
            <div style="color: #94a3b8; font-size: 0.85rem;">severity levels</div>
        </div>
        """, unsafe_allow_html=True)
    
    with spec_col3:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Device</div>
            <div class="metric-value">{device.type.upper()}</div>
            <div style="color: #94a3b8; font-size: 0.85rem;">compute</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="info-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **💾 Model Checkpoint:**
    - Path: `{MODEL_PATH}`
    - Status: {'✅ Loaded Successfully' if os.path.exists(MODEL_PATH) else '⚠️ Not Found'}
    - Framework: PyTorch
    - Architecture: Dual-Transformer Fusion
    """)
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.page == 'report':
    # Back button
    col_back, col_spacer = st.columns([1, 5])
    with col_back:
        if st.button("← Back", key="back_to_home", use_container_width=True):
            st.session_state.page = 'home'
            st.rerun()
    
    st.markdown("""
    <div class="main-header">
        <h1 class="main-title">📑 Detailed Analysis Report</h1>
        <p class="subtitle">AI-Generated Environmental Assessment</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Report content
    #st.markdown('<div class="custom-card" style="min-height: 60vh;">', unsafe_allow_html=True)
    
    if 'generated_report' in st.session_state:
        st.markdown(st.session_state['generated_report'])
    else:
        st.warning("No report available. Please run an analysis first.")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Download button positioned at bottom right
    if 'generated_report' in st.session_state:
        st.markdown('<div style="margin-top: 2rem; display: flex; justify-content: flex-end;">', unsafe_allow_html=True)
        
        # Get stored prediction data
        pred_name = st.session_state.get('pred_name', 'Unknown')
        pred_prob = st.session_state.get('pred_prob', 0.0)
        salinity = st.session_state.get('salinity', 0.0)
        water_temp = st.session_state.get('water_temp', 0.0)
        cellcount = st.session_state.get('cellcount', 0.0)
        
        # Create Word document
        word_buffer = create_word_report(
            st.session_state['generated_report'],
            pred_name,
            pred_prob,
            salinity,
            water_temp,
            cellcount
        )
        
        st.download_button(
            label="📥 Download Report as Word",
            data=word_buffer,
            file_name=f"HAB_Report_{time.strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_report",
            use_container_width=False
        )
        
        st.markdown('</div>', unsafe_allow_html=True)
